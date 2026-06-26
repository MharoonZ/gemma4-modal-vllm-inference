from __future__ import annotations

import math
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_DOCX = Path("Gemma4_Modal_vLLM_Architecture_Explained.docx")
DIAGRAM_DIR = Path("architecture_diagrams")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    max_width: int,
    text_font: ImageFont.ImageFont,
    fill: str = "#17202a",
    line_gap: int = 4,
) -> int:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if draw.textbbox((0, 0), candidate, font=text_font)[2] <= max_width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)

    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=text_font, fill=fill)
        y += text_font.size + line_gap
    return y


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str = "#ffffff",
    outline: str = "#cfd7e6",
    title_fill: str = "#0f172a",
) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, _ = box
    draw_wrapped_text(draw, title, (x1 + 22, y1 + 18), x2 - x1 - 44, font(24, True), fill=title_fill)
    draw_wrapped_text(draw, body, (x1 + 22, y1 + 58), x2 - x1 - 44, font(18), fill="#475467")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#0f766e",
    width: int = 5,
) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=color)


def new_canvas(width: int = 1800, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "#f6f8fb")
    return img, ImageDraw.Draw(img)


def save_architecture_diagram() -> Path:
    img, draw = new_canvas(1800, 950)
    draw.text((60, 45), "Current System Architecture", font=font(36, True), fill="#0f172a")
    draw.text(
        (60, 95),
        "A custom UI calls our own Modal-hosted vLLM server, which runs Gemma 4 on a GPU.",
        font=font(22),
        fill="#475467",
    )

    boxes = {
        "user": (70, 220, 360, 420),
        "ui": (480, 220, 790, 420),
        "modal": (910, 220, 1220, 420),
        "vllm": (1340, 220, 1660, 420),
        "gpu": (1010, 590, 1320, 790),
        "cache": (1430, 590, 1700, 790),
    }
    rounded_box(draw, boxes["user"], "User", "Types a prompt and reads the answer in the browser.")
    rounded_box(draw, boxes["ui"], "Micro UI", "Local HTML page with prompt input, send button, latency log, and token usage.")
    rounded_box(draw, boxes["modal"], "Modal Endpoint", "Public HTTPS endpoint that routes requests to the running container.")
    rounded_box(draw, boxes["vllm"], "vLLM Server", "OpenAI-compatible API that loads, schedules, and serves Gemma 4.")
    rounded_box(draw, boxes["gpu"], "GPU Container", "L40S for prototype. H100 remains fallback for heavier loads.")
    rounded_box(draw, boxes["cache"], "Model Caches", "Hugging Face and vLLM volumes reduce repeat download and compile work.")

    arrow(draw, (360, 320), (480, 320))
    arrow(draw, (790, 320), (910, 320))
    arrow(draw, (1220, 320), (1340, 320))
    arrow(draw, (1500, 420), (1210, 590))
    arrow(draw, (1500, 420), (1565, 590))
    draw.text((405, 285), "Prompt", font=font(18, True), fill="#0f766e")
    draw.text((810, 285), "HTTPS", font=font(18, True), fill="#0f766e")
    draw.text((1245, 285), "/v1/chat/completions", font=font(18, True), fill="#0f766e")

    path = DIAGRAM_DIR / "01_current_architecture.png"
    img.save(path)
    return path


def save_comparison_diagram() -> Path:
    img, draw = new_canvas(1800, 1000)
    draw.text((60, 45), "Direct Chatbot vs Our Custom Stack", font=font(36, True), fill="#0f172a")

    draw.text((170, 130), "Using ChatGPT/Gemini Directly", font=font(28, True), fill="#0f172a")
    draw.text((1070, 130), "Using Our Gemma 4 Stack", font=font(28, True), fill="#0f172a")

    left = [
        ((90, 210, 390, 360), "User", "Uses provider's app or API."),
        ((520, 210, 820, 360), "Provider Platform", "Provider controls UI, serving, safety, logs, models, GPUs."),
        ((300, 520, 610, 700), "Hosted Model", "Model runs inside provider infrastructure."),
    ]
    right = [
        ((980, 210, 1280, 360), "User", "Uses our micro UI."),
        ((1410, 210, 1710, 360), "Our UI + Endpoint", "We control request flow, metrics, and product behavior."),
        ((980, 520, 1280, 700), "vLLM + Gemma 4", "Open model served in our Modal app."),
        ((1410, 520, 1710, 700), "Modal GPU", "Cloud GPU runtime billed by seconds."),
    ]
    for box, title, body in left + right:
        rounded_box(draw, box, title, body)
    arrow(draw, (390, 285), (520, 285))
    arrow(draw, (670, 360), (470, 520))
    arrow(draw, (1280, 285), (1410, 285))
    arrow(draw, (1560, 360), (1130, 520))
    arrow(draw, (1280, 610), (1410, 610))

    draw.rounded_rectangle((80, 790, 830, 920), radius=18, fill="#fff7ed", outline="#fed7aa", width=3)
    draw_wrapped_text(
        draw,
        "Main idea: direct chatbot products hide almost all infrastructure. You use the finished service.",
        (110, 830),
        690,
        font(22, True),
        fill="#9a3412",
    )
    draw.rounded_rectangle((960, 790, 1720, 920), radius=18, fill="#ecfdf5", outline="#a7f3d0", width=3)
    draw_wrapped_text(
        draw,
        "Main idea: our stack gives us control over the UI, model choice, endpoint behavior, latency tracking, and deployment path.",
        (990, 830),
        700,
        font(22, True),
        fill="#065f46",
    )
    path = DIAGRAM_DIR / "02_direct_vs_custom.png"
    img.save(path)
    return path


def save_sequence_diagram() -> Path:
    img, draw = new_canvas(1800, 1050)
    draw.text((60, 45), "Prompt Request Lifecycle", font=font(36, True), fill="#0f172a")
    steps = [
        ("1. User enters prompt", "Example: 'Can you extract info from pdf'."),
        ("2. UI checks health", "Browser calls GET /health before sending the prompt."),
        ("3. UI sends chat request", "Browser posts JSON to /v1/chat/completions."),
        ("4. Modal routes request", "If no container is warm, Modal starts one on GPU."),
        ("5. vLLM runs inference", "Prompt becomes tokens, Gemma 4 generates response tokens."),
        ("6. UI shows result", "Response, latency, health wait, and token count are displayed."),
    ]
    x = 90
    y = 170
    box_w = 500
    box_h = 150
    positions = []
    for i, (title, body) in enumerate(steps):
        row = i // 3
        col = i % 3
        bx = x + col * 570
        by = y + row * 300
        positions.append((bx, by, bx + box_w, by + box_h))
        rounded_box(draw, positions[-1], title, body)
    for i in range(2):
        arrow(draw, (positions[i][2], positions[i][1] + 75), (positions[i + 1][0], positions[i + 1][1] + 75))
    arrow(draw, (positions[2][0] + 250, positions[2][3]), (positions[5][0] + 250, positions[5][1]))
    arrow(draw, (positions[5][0], positions[5][1] + 75), (positions[4][2], positions[4][1] + 75))
    arrow(draw, (positions[4][0], positions[4][1] + 75), (positions[3][2], positions[3][1] + 75))

    draw.rounded_rectangle((150, 810, 1650, 970), radius=18, fill="#eff6ff", outline="#bfdbfe", width=3)
    draw_wrapped_text(
        draw,
        "Cold start happens when Modal must start a GPU container, load model weights, and let vLLM warm up. Warm request happens when the container is already running.",
        (190, 855),
        1420,
        font(24, True),
        fill="#1e3a8a",
    )
    path = DIAGRAM_DIR / "03_request_lifecycle.png"
    img.save(path)
    return path


def save_cost_diagram() -> Path:
    img, draw = new_canvas(1800, 950)
    draw.text((60, 45), "Cost Model: Token Billing vs GPU Runtime Billing", font=font(36, True), fill="#0f172a")
    rounded_box(
        draw,
        (110, 180, 790, 430),
        "Direct API Products",
        "Often priced per input and output token, or bundled into a subscription. The provider owns the GPU runtime.",
        fill="#fff7ed",
        outline="#fed7aa",
    )
    rounded_box(
        draw,
        (1010, 180, 1690, 430),
        "Our Modal/vLLM Prototype",
        "Modal charges for GPU seconds while the container is starting, loading, warming up, or serving requests.",
        fill="#ecfdf5",
        outline="#a7f3d0",
    )
    arrow(draw, (790, 305), (1010, 305), color="#64748b")
    draw.text((850, 270), "Different billing model", font=font(20, True), fill="#334155")

    bullets = [
        ("Model cold start", "Can cost more than a tiny prompt because the GPU is reserved while the model loads."),
        ("Warm endpoint", "Useful for low latency, but it costs money while kept alive."),
        ("Token usage", "Still important for capacity planning, but not the main Modal charge in this prototype."),
        ("L40S", "Validated for this prototype and cheaper than H100."),
    ]
    y = 550
    for title, body in bullets:
        draw.rounded_rectangle((180, y, 1620, y + 70), radius=14, fill="#ffffff", outline="#cfd7e6", width=2)
        draw.text((220, y + 18), title, font=font(20, True), fill="#0f172a")
        draw_wrapped_text(draw, body, (490, y + 18), 1070, font(18), fill="#475467")
        y += 88
    path = DIAGRAM_DIR / "04_cost_model.png"
    img.save(path)
    return path


def add_hyperlink(paragraph, text: str, url: str) -> None:
    # Keep this public-API only so the generator remains compatible across
    # python-docx versions. Word will still auto-detect the visible URL.
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(5, 99, 193)
    run.underline = True


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    for line in text.splitlines():
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def build_document(diagrams: dict[str, Path]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(14)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Gemma 4 + vLLM + Modal Micro-UI Architecture")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plain-English explanation for someone used to ChatGPT/Gemini direct usage")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 84, 103)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared: April 30, 2026").italic = True

    document.add_paragraph()
    document.add_picture(str(diagrams["architecture"]), width=Inches(7.0))

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "We are building a custom AI application stack instead of only using a finished hosted chatbot product like ChatGPT or Gemini. "
        "The current prototype has a browser micro-UI that sends prompts to a Modal-hosted vLLM endpoint. vLLM then serves Google's Gemma 4 open model on a cloud GPU and returns the answer to the UI."
    )
    document.add_paragraph(
        "The important shift is this: with ChatGPT/Gemini direct usage, the provider owns the complete product and infrastructure. "
        "With this stack, we choose the model, host it through our own endpoint, measure latency and tokens ourselves, and can build product-specific features on top."
    )
    add_bullets(
        document,
        [
            "Current model: google/gemma-4-E4B-it.",
            "Serving engine: vLLM 0.19.0 with an OpenAI-compatible chat completion API.",
            "Cloud runtime: Modal GPU container, validated on L40S for the prototype.",
            "User interface: local browser micro-UI with prompt input, response display, health checks, latency log, and token usage.",
            "Current status: the model loads, endpoint responds, UI works, and L40S is sufficient for this prototype stage.",
        ],
    )

    document.add_heading("2. The Simplest Mental Model", level=1)
    document.add_paragraph(
        "Think of a normal chatbot as a finished restaurant: you walk in, order food, and the restaurant hides the kitchen, suppliers, staff, equipment, and recipes. "
        "Using ChatGPT or Gemini directly is like that. You type in a prompt, and the provider handles everything behind the scenes."
    )
    document.add_paragraph(
        "Our system is more like building our own small kitchen. We still use cloud equipment, but we decide which model to run, how to expose it, how to measure it, and what application experience to build around it."
    )
    document.add_paragraph("In one line:")
    add_bullets(
        document,
        [
            "Traditional use: User -> ChatGPT/Gemini app -> provider-owned model.",
            "Our system: User -> our micro-UI -> our Modal endpoint -> vLLM -> Gemma 4 running on GPU.",
        ],
    )

    document.add_heading("3. What We Have Built So Far", level=1)
    document.add_paragraph(
        "The current project is intentionally small. It is not yet a full product. It is a working foundation that proves we can run an open model and build a UI on top of it."
    )
    add_table(
        document,
        ["Part", "What it does"],
        [
            ["micro_ui.html", "A browser page where the user enters a prompt and sees the answer."],
            ["gemma4_modal_vllm_app.py", "The Modal app that starts vLLM and serves Gemma 4."],
            ["start_gemma4_micro_ui.ps1", "Starts the local UI server and starts the Modal development endpoint."],
            ["run_gemma4_smoke_test.ps1", "Runs a command-line Hello World test without the UI."],
            ["MICRO_UI_IMPLEMENTATION.md", "Operational notes, test results, and run instructions."],
        ],
    )
    document.add_paragraph(
        "The UI currently supports endpoint input, prompt input, model name, max token setting, timeout setting, health check, send button, loading/ready state, response display, first-run latency, warm latency, health wait timing, token usage, and request history."
    )

    document.add_heading("4. Architecture Diagram", level=1)
    document.add_paragraph(
        "This diagram shows the main moving parts. The user does not call Gemma 4 directly. The browser calls the Modal endpoint, Modal routes traffic to the vLLM server, and vLLM runs Gemma 4 on the GPU."
    )
    document.add_picture(str(diagrams["architecture"]), width=Inches(7.2))

    document.add_heading("5. How This Differs From ChatGPT or Gemini Direct Use", level=1)
    document.add_picture(str(diagrams["comparison"]), width=Inches(7.2))
    add_table(
        document,
        ["Area", "ChatGPT/Gemini Directly", "Our Gemma 4 + vLLM + Modal Stack"],
        [
            ["Who owns the model runtime?", "The provider owns and operates it.", "We run an open model in our own Modal app."],
            ["Model choice", "Limited to what the provider exposes.", "We can choose Gemma 4 now, and later test other open models."],
            ["UI control", "Mostly fixed by the provider app.", "We design the UI and workflow ourselves."],
            ["API shape", "Provider-specific API or app.", "OpenAI-compatible vLLM endpoint, currently /v1/chat/completions."],
            ["Cost model", "Often subscription or per-token API pricing.", "GPU runtime billing on Modal; token count is still measured but not the primary Modal charge."],
            ["Customization", "Prompting and settings only, unless using provider tools.", "We can modify serving settings, routing, UI behavior, logging, and future file workflows."],
            ["Operational responsibility", "Provider handles infrastructure.", "We must manage endpoint lifecycle, cost, errors, deployment, and security."],
            ["Data/control posture", "Data goes to provider service.", "Data goes through our app and cloud runtime, giving more control over handling and logging choices."],
        ],
    )

    document.add_heading("6. Component-by-Component Explanation", level=1)
    document.add_heading("6.1 Micro UI", level=2)
    document.add_paragraph(
        "The micro-UI is the small web page we created. It is the visible layer. A user enters a prompt, clicks Send, and sees the model response. It also shows operational details that normal chatbot products hide, such as health wait, first-run latency, warm latency, and token usage."
    )
    document.add_heading("6.2 Modal Endpoint", level=2)
    document.add_paragraph(
        "Modal provides the public HTTPS endpoint. When the UI sends a request to the endpoint, Modal routes that request to the cloud container running our vLLM server. If the container is not running, Modal may need to start one. That is called a cold start."
    )
    document.add_heading("6.3 vLLM", level=2)
    document.add_paragraph(
        "vLLM is not the AI model. It is the serving engine. Its job is to load the model, expose API routes, tokenize prompts, batch/schedule requests efficiently, run inference on the GPU, and return responses in an OpenAI-compatible format."
    )
    document.add_heading("6.4 Gemma 4", level=2)
    document.add_paragraph(
        "Gemma 4 is the open model family we started with. In this prototype, the selected model is google/gemma-4-E4B-it. It is instruction-tuned, meaning it is meant to respond to user instructions and chat-style prompts."
    )
    document.add_heading("6.5 GPU Container", level=2)
    document.add_paragraph(
        "Large language models need GPU memory and compute to respond quickly. The model is loaded into a Modal GPU container. We first validated on H100, then tested L40S and confirmed L40S works for this prototype. That matters because L40S is lower cost than H100."
    )
    document.add_heading("6.6 Hugging Face and vLLM Caches", level=2)
    document.add_paragraph(
        "The app mounts cache volumes for Hugging Face and vLLM. These caches reduce repeated work. The first run may download model weights and compile runtime graphs. Later runs can reuse cached assets, reducing startup cost and latency."
    )

    document.add_heading("7. What Happens When a User Sends a Prompt", level=1)
    document.add_picture(str(diagrams["sequence"]), width=Inches(7.2))
    add_numbered(
        document,
        [
            "The user types a prompt into the micro-UI.",
            "The UI checks the endpoint health by calling GET /health.",
            "If healthy, the UI sends JSON to POST /v1/chat/completions.",
            "Modal routes the request to the vLLM server running in the GPU container.",
            "vLLM tokenizes the prompt and schedules the request.",
            "Gemma 4 generates output tokens on the GPU.",
            "vLLM returns the response in OpenAI-compatible JSON.",
            "The UI displays the text response and records latency/token metrics.",
        ],
    )

    document.add_heading("8. API Shape Used by the UI", level=1)
    document.add_paragraph(
        "The micro-UI sends requests using the same general shape as OpenAI-compatible chat completion APIs. That is useful because it keeps the application interface familiar and makes it easier to swap models or serving backends later."
    )
    add_code_block(
        document,
        """POST /v1/chat/completions
{
  "model": "google/gemma-4-E4B-it",
  "messages": [
    {"role": "user", "content": "Hello World"}
  ],
  "max_tokens": 128,
  "temperature": 0,
  "stream": false,
  "chat_template_kwargs": {"enable_thinking": false}
}""",
    )
    document.add_paragraph(
        "The response includes the assistant message and usage metadata. The UI reads choices[0].message.content for the answer and usage.total_tokens for token usage."
    )

    document.add_heading("9. Tokens Explained Simply", level=1)
    document.add_paragraph(
        "A token is a small piece of text. It can be a word, part of a word, punctuation, or spacing. Models do not read text exactly as humans do. They read and generate tokens."
    )
    add_bullets(
        document,
        [
            "Prompt tokens are tokens sent into the model.",
            "Completion tokens are tokens generated by the model.",
            "Total tokens are prompt tokens plus completion tokens.",
            "Max tokens limits how many output tokens the model may generate.",
        ],
    )
    document.add_paragraph(
        "In direct API products, token count often directly affects the bill. In this Modal prototype, token usage is still important, but Modal's main cost is GPU runtime seconds. Token usage helps us understand workload size, capacity, and future cost if we later move to token-priced services or add internal accounting."
    )

    document.add_heading("10. Latency Terms Explained", level=1)
    add_table(
        document,
        ["Term", "Meaning"],
        [
            ["Health wait", "How long the UI waited for /health to confirm the endpoint is alive."],
            ["First-run latency", "The first successful request in a browser session. It may include cold-start effects."],
            ["Warm latency", "Latency after the endpoint is already running and model is already loaded."],
            ["Request latency", "Time for the actual chat completion request after health is ready."],
            ["Cold start", "Starting a GPU container, loading the model, compiling/warming vLLM, then becoming ready."],
            ["Warm request", "A request handled while the GPU container and model are already ready."],
        ],
    )

    document.add_heading("11. Cost Model and What We Learned", level=1)
    document.add_picture(str(diagrams["cost"]), width=Inches(7.2))
    document.add_paragraph(
        "The credits used during testing were mostly from GPU runtime, not prompt tokens. Modal charges while the GPU container is starting, loading the model, warming up, serving requests, or being kept alive for dev testing."
    )
    add_table(
        document,
        ["Observed Item", "Result"],
        [
            ["Successful L40S validation", "Passed. Request latency about 2.468 seconds after readiness."],
            ["Model loading memory on L40S", "About 14.25 GiB."],
            ["Health/startup wait on L40S validation", "About 158.664 seconds."],
            ["Endpoint", "https://foodqer--gemma4-vllm-micro-ui-serve-dev.modal.run"],
            ["Recommendation", "Use L40S for prototype; keep H100 only as fallback for larger context or heavier load."],
        ],
    )
    document.add_paragraph(
        "The earlier approximately $2.99 credit usage was explained by model cold starts, vLLM compile/warmup, dev endpoint runtime, and one failed/timeout attempt while fixing browser/CORS behavior. The prompt token usage itself was small."
    )

    document.add_heading("12. What This System Is Not Yet", level=1)
    document.add_paragraph(
        "This prototype proves the serving path and UI path. It is not yet a finished production chatbot or document processing system."
    )
    add_bullets(
        document,
        [
            "It does not yet upload or parse PDF files directly.",
            "It does not yet store long-term conversation history on a backend database.",
            "It does not yet have authentication or user accounts.",
            "It does not yet have production monitoring, alerts, or rate limits.",
            "It currently uses a development Modal endpoint, not a hardened production deployment.",
        ],
    )

    document.add_heading("13. Why This Architecture Is Useful", level=1)
    add_bullets(
        document,
        [
            "Model flexibility: we can start with Gemma 4 and later test other open models.",
            "Product control: we can build a workflow-specific UI instead of using a generic chatbot interface.",
            "Observability: the UI already tracks latency and token usage.",
            "Cost experimentation: we validated L40S as a lower-cost GPU path.",
            "OpenAI-compatible API: the application layer can remain familiar even if the backend model changes.",
            "Incremental development: we can add PDF upload, retrieval, formatting, or structured outputs step by step.",
        ],
    )

    document.add_heading("14. Operational Rules Going Forward", level=1)
    add_bullets(
        document,
        [
            "Use L40S for prototype testing unless a larger model or context requires H100.",
            "Stop the Modal dev endpoint when testing is complete.",
            "Check active containers with modal container list if costs look unexpected.",
            "Keep scaledown/min-containers conservative so idle GPUs do not keep running.",
            "Add a Hugging Face token secret only if downloads are rate-limited or slow.",
            "Move from modal serve to a controlled deployment path when the UI is ready for broader testing.",
        ],
    )

    document.add_heading("15. Next Steps", level=1)
    add_numbered(
        document,
        [
            "Add file upload support for PDFs or text files.",
            "Parse PDF text before sending content to the model.",
            "Add a backend proxy so the browser does not directly manage the Modal endpoint.",
            "Add authentication before exposing the UI to real users.",
            "Add request logging, cost tracking, and error reporting.",
            "Improve output formatting for summaries, extracted fields, tables, and citations.",
            "Test other open models and compare quality, latency, memory, and cost.",
        ],
    )

    document.add_heading("16. Glossary", level=1)
    add_table(
        document,
        ["Term", "Plain-English Meaning"],
        [
            ["LLM", "Large language model. The AI model that reads prompts and generates text."],
            ["Open model", "A model whose weights or usage path are available outside a closed provider product."],
            ["Gemma 4", "The open model family being tested first."],
            ["vLLM", "The server software that runs the model efficiently and exposes API routes."],
            ["Modal", "The cloud platform that runs our GPU container and provides the web endpoint."],
            ["GPU", "Specialized hardware used to run model inference efficiently."],
            ["Endpoint", "A web URL that software can call to send requests."],
            ["Cold start", "Startup work before the model is ready."],
            ["Warm request", "A request served after the model is already running."],
            ["Token", "A chunk of text used by the model."],
            ["CORS", "Browser security rule controlling whether a web page can call another domain."],
            ["Inference", "The act of running the model to generate an answer."],
        ],
    )

    document.add_heading("17. Final Summary", level=1)
    document.add_paragraph(
        "If you have only used ChatGPT or Gemini directly, the key thing to understand is that we are no longer only using a finished chatbot. "
        "We are assembling the pieces of our own AI application: UI, endpoint, serving engine, open model, GPU runtime, metrics, and future workflow features."
    )
    document.add_paragraph(
        "This gives more control and flexibility, but it also creates engineering responsibilities around deployment, latency, cost, reliability, and security. "
        "The current prototype proves the foundation and gives us a path to build product-specific functionality in small, safe steps."
    )

    document.add_heading("References", level=1)
    for label, url in [
        ("Modal pricing", "https://modal.com/pricing"),
        ("Modal vLLM example", "https://modal.com/docs/examples/vllm_inference"),
        ("vLLM Gemma 4 recipe", "https://docs.vllm.ai/projects/recipes/en/latest/Google/Gemma4.html"),
        ("Hugging Face Gemma 4 E4B model card", "https://huggingface.co/google/gemma-4-E4B-it"),
    ]:
        p = document.add_paragraph()
        p.add_run(label + ": ")
        add_hyperlink(p, url, url)

    document.save(OUTPUT_DOCX)


def main() -> None:
    DIAGRAM_DIR.mkdir(exist_ok=True)
    diagrams = {
        "architecture": save_architecture_diagram(),
        "comparison": save_comparison_diagram(),
        "sequence": save_sequence_diagram(),
        "cost": save_cost_diagram(),
    }
    build_document(diagrams)
    print(OUTPUT_DOCX.resolve())


if __name__ == "__main__":
    main()
